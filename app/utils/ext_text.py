import fitz
import docx
import io
from io import BytesIO


def extracted_text_from_pdf(resume_file):
    try:
        # Handle both FileStorage and BytesIO
        if hasattr(resume_file, "read"):
            file_bytes = resume_file.read()
            resume_file.seek(0)
        else:
            raise ValueError("Invalid file object for PDF parsing")

        text = ""
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        print("PDF parsing error:", e)
        return {"error": f"Failed to read PDF: {str(e)}"}


def extracted_text_from_docx(resume_file):
    try:
        # If FileStorage (Flask upload), use .stream
        if hasattr(resume_file, "stream"):
            document = docx.Document(resume_file.stream)
        else:
            # If BytesIO (tests), use directly
            document = docx.Document(resume_file)

        text = ""
        for para in document.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        print("DOCX parsing error:", e)
        return {"error": f"Failed to read DOCX: {str(e)}"}
